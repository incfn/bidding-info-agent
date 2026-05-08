# -*- coding: utf-8 -*-
"""
招标信息智能分析 Agent
读取 ccgp_results.csv，调用 DeepSeek 大模型逐条分析，
输出结构化 CSV 和 Word 简报。
"""

import csv
import json
import os
import re
import time
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from openai import OpenAI


def load_env():
    """从 .env 文件加载环境变量"""
    env_path = os.path.join(os.path.dirname(__file__), ".env")
    if os.path.exists(env_path):
        with open(env_path, "r", encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if not line or line.startswith("#"):
                    continue
                if "=" in line:
                    key, value = line.split("=", 1)
                    os.environ.setdefault(key.strip(), value.strip())


load_env()

# ========================= 配置区域 =========================

# DeepSeek API 配置
# 方式1：在 .env 文件中设置 DEEPSEEK_API_KEY（推荐）
# 方式2：设置系统环境变量 DEEPSEEK_API_KEY
DEEPSEEK_API_KEY = os.getenv("DEEPSEEK_API_KEY", "")
DEEPSEEK_BASE_URL = "https://api.deepseek.com"
DEEPSEEK_MODEL = "deepseek-chat"  # 分析任务用 chat 足够，便宜且快

# 输入输出文件
INPUT_CSV = "ccgp_results.csv"
OUTPUT_CSV = "ccgp_analyzed.csv"
OUTPUT_DOCX = "招标情报分析简报.docx"

# 进度缓存文件（中断后可续跑）
PROGRESS_FILE = "_analysis_progress.json"

# 每次 API 调用间隔（秒），避免请求过快
API_DELAY = 0.5

# ========================= Prompt 设计 =========================

SYSTEM_PROMPT = """你是一位医疗器械行业的市场情报分析师，目前在施耐德电气销售部工作。
请根据提供的招标公告信息，进行结构化分析和相关度评估，只输出纯 JSON，不要有任何 markdown 代码块标记或其他解释性文字。"""

USER_PROMPT_TEMPLATE = """【招标公告信息】
标题：{标题}
发布时间：{发布时间}
项目编号：{项目编号}
项目名称：{项目名称}
预算金额：{预算金额}
采购需求：{采购需求}
提交投标文件截止时间：{提交投标文件截止时间}
项目联系人：{项目联系人}
电话：{电话}
采购人名称：{采购人名称}
采购人地址：{采购人地址}
详情页链接：{详情页链接}

请对以上招标公告进行深度分析，输出以下严格的 JSON 格式：
{{
    "相关度评分": 1-10,
    "判断理由": "为什么相关或不相关，50字以内",
    "关键信息": {{
        "地区": "项目所在省市，从地址或标题推断",
        "项目类型": "设备采购/工程建设/信息化/耗材采购/服务/其他",
        "预算规模": "大额(>500万)/中等(100-500万)/小额(<100万)/未知",
        "核心产品": "招标中涉及的关键设备或系统，没有则填无",
        "潜在机会": "施耐德可以切入的具体环节"
    }},
    "跟进建议": "销售团队下一步可以怎么做，30字以内"
}}

评分规则（严格执行）：
- 10分：明确的大型医疗设备/配电系统/实验室设备/医院基建采购
- 7-9分：医药产业园建设、检验检测机构设备采购、信息化改造涉及电气
- 4-6分：医药相关但机会不明确，或预算很小
- 1-3分：完全不相关（如餐饮服务、班车租赁、普通试剂耗材配送、办公用品等）

注意：
1. 只返回 JSON，不要 ```json 代码块。
2. 如果字段为空，请基于标题和采购需求合理推断。
3. 地区尽量精确到省市。
"""

# ========================= 核心函数 =========================

def get_deepseek_client():
    """初始化 DeepSeek 客户端"""
    if not DEEPSEEK_API_KEY or DEEPSEEK_API_KEY == "your_api_key_here":
        raise ValueError(
            "[ERROR] 未配置 DeepSeek API Key！\n"
            "请在系统环境变量中设置 DEEPSEEK_API_KEY，\n"
            "或直接修改脚本中的 DEEPSEEK_API_KEY 变量。"
        )
    return OpenAI(api_key=DEEPSEEK_API_KEY, base_url=DEEPSEEK_BASE_URL)


def call_deepseek(client, prompt):
    """调用 DeepSeek API，返回解析后的 JSON 字典"""
    try:
        response = client.chat.completions.create(
            model=DEEPSEEK_MODEL,
            messages=[
                {"role": "system", "content": SYSTEM_PROMPT},
                {"role": "user", "content": prompt},
            ],
            temperature=0.3,  # 低温度，保证输出稳定
            max_tokens=800,
        )
        content = response.choices[0].message.content.strip()

        # 有时候模型会包一层 ```json ... ```，这里做兜底清理
        if content.startswith("```"):
            content = content.strip("`").strip()
            if content.lower().startswith("json"):
                content = content[4:].strip()

        return json.loads(content)
    except Exception as e:
        print(f"      [!] API 调用或解析失败: {e}")
        return None


def set_run_font(run, font_name='微软雅黑', font_size=None, bold=False):
    """强制统一设置 run 的字体（同时覆盖西文和中文）"""
    from docx.oxml.ns import qn
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if font_size:
        run.font.size = Pt(font_size)
    if bold:
        run.font.bold = True


def clean_pub_time(val):
    """发布时间只保留年月日"""
    val = str(val).strip()
    if " " in val:
        return val.split(" ")[0].strip()
    return val


def clean_budget(val):
    """预算金额统一单位并加千位符"""
    val = str(val).strip().strip('"').strip("'")
    if not val or val in ("详见公告正文", "无", "0", "-"):
        return val

    # 如果已经清洗过，直接返回
    if val.startswith("¥") and ("万元" in val or "元" in val):
        return val

    # 去掉货币符号、"人民币"字样及残留括号
    val_clean = val.replace("￥", "").replace("¥", "").replace("人民币", "").strip()
    val_clean = val_clean.replace("（", "").replace("）", "").replace("(", "").replace(")", "").strip()

    if "万元" in val_clean:
        num_str = val_clean.replace("万元", "").strip()
        try:
            num = float(num_str.replace(",", ""))
            return f"¥{num:,.2f} 万元"
        except ValueError:
            return val
    elif "元" in val_clean:
        num_str = val_clean.replace("元", "").strip()
        try:
            num = float(num_str.replace(",", ""))
            if num >= 10000:
                return f"¥{num / 10000:,.2f} 万元"
            else:
                return f"¥{num:,.2f} 元"
        except ValueError:
            return val
    else:
        # 纯数字，尝试解析
        try:
            num = float(val_clean.replace(",", ""))
            if num >= 10000:
                return f"¥{num / 10000:,.2f} 万元"
            else:
                return f"¥{num:,.2f} 元"
        except ValueError:
            return val


def clean_phone(val):
    """电话统一用 / 分隔"""
    val = str(val).strip()
    if not val or val in ("-", "无"):
        return val

    # 统一分隔符
    val = re.sub(r"[、，,\s/]+", "/", val)
    val = re.sub(r"/+", "/", val)
    val = val.strip("/")
    phones = [p.strip() for p in val.split("/") if p.strip()]
    return " / ".join(phones)


def clean_row(row):
    """清洗单条数据"""
    row = dict(row)
    row["发布时间"] = clean_pub_time(row.get("发布时间", ""))
    row["预算金额"] = clean_budget(row.get("预算金额", ""))
    row["电话"] = clean_phone(row.get("电话", ""))
    return row


def load_progress():
    """加载已分析的进度"""
    if os.path.exists(PROGRESS_FILE):
        with open(PROGRESS_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    return {}


def save_progress(progress):
    """保存分析进度"""
    with open(PROGRESS_FILE, "w", encoding="utf-8") as f:
        json.dump(progress, f, ensure_ascii=False, indent=2)


def analyze_csv():
    """主流程：读取 CSV → 调用 DeepSeek → 写入 OUTPUT_CSV"""
    client = get_deepseek_client()
    progress = load_progress()

    # 清洗已有缓存中的数据
    if progress:
        progress = {k: clean_row(v) for k, v in progress.items()}
        save_progress(progress)

    # 读取原始数据
    with open(INPUT_CSV, "r", encoding="utf-8-sig") as f:
        reader = csv.DictReader(f)
        rows = list(reader)

    total = len(rows)
    print(f"[*] 共读取到 {total} 条招标信息")
    print(f"[*] 已分析进度: {len(progress)} 条")

    # 准备输出字段
    base_fields = list(rows[0].keys()) if rows else []
    extra_fields = [
        "相关度评分",
        "判断理由",
        "地区",
        "项目类型",
        "预算规模",
        "核心产品",
        "潜在机会",
        "跟进建议",
    ]
    output_fields = base_fields + extra_fields

    # 先写出已有结果（避免中断丢失）
    with open(OUTPUT_CSV, "w", newline="", encoding="utf-8-sig") as f:
        writer = csv.DictWriter(f, fieldnames=output_fields)
        writer.writeheader()

    for idx, row in enumerate(rows, 1):
        row_id = str(idx)

        # 如果已经分析过，直接复用（已清洗）
        if row_id in progress:
            analyzed = progress[row_id]
            print(f"  [{idx}/{total}] [OK] 已分析(缓存): {row.get('标题', '')[:40]}...")
        else:
            print(f"  [{idx}/{total}] [分析中] {row.get('标题', '')[:40]}...")

            prompt = USER_PROMPT_TEMPLATE.format(**row)
            result = call_deepseek(client, prompt)

            if result is None:
                # API 失败，填充空值跳过
                result = {
                    "相关度评分": "",
                    "判断理由": "分析失败",
                    "关键信息": {
                        "地区": "",
                        "项目类型": "",
                        "预算规模": "",
                        "核心产品": "",
                        "潜在机会": "",
                    },
                    "跟进建议": "",
                }

            # 扁平化结果
            analyzed = dict(row)
            analyzed["相关度评分"] = result.get("相关度评分", "")
            analyzed["判断理由"] = result.get("判断理由", "")
            key_info = result.get("关键信息", {})
            analyzed["地区"] = key_info.get("地区", "")
            analyzed["项目类型"] = key_info.get("项目类型", "")
            analyzed["预算规模"] = key_info.get("预算规模", "")
            analyzed["核心产品"] = key_info.get("核心产品", "")
            analyzed["潜在机会"] = key_info.get("潜在机会", "")
            analyzed["跟进建议"] = result.get("跟进建议", "")

            # 清洗后存入缓存
            analyzed = clean_row(analyzed)
            progress[row_id] = analyzed
            save_progress(progress)
            time.sleep(API_DELAY)

        # 追加写入 CSV
        with open(OUTPUT_CSV, "a", newline="", encoding="utf-8-sig") as f:
            writer = csv.DictWriter(f, fieldnames=output_fields)
            writer.writerow(analyzed)

    print(f"\n[√] 分析完成，结果已保存到: {OUTPUT_CSV}")
    return list(progress.values())


def generate_word_report(rows):
    """生成 Word 简报"""
    doc = Document()

    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("医疗器械招标情报分析简报")
    set_run_font(title_run, font_name='微软雅黑', font_size=18, bold=True)

    # 日期
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f"生成时间：{datetime.now().strftime('%Y年%m月%d日')}")
    set_run_font(date_run, font_name='微软雅黑', font_size=10.5)

    doc.add_paragraph()

    # 按评分分组
    high = []   # >= 8
    medium = [] # 6-7
    low = []    # <= 5 or 空
    failed = []

    for r in rows:
        score = r.get("相关度评分", "")
        try:
            s = float(score)
            if s >= 8:
                high.append(r)
            elif s >= 6:
                medium.append(r)
            else:
                low.append(r)
        except (ValueError, TypeError):
            if r.get("判断理由") == "分析失败":
                failed.append(r)
            else:
                low.append(r)

    def add_section(heading_text, items, color_rgb=(0, 0, 0)):
        if not items:
            return
        h = doc.add_heading(heading_text, level=1)
        set_run_font(h.runs[0], font_name='微软雅黑', font_size=14, bold=True)
        h.runs[0].font.color.rgb = RGBColor(*color_rgb)

        for i, item in enumerate(items, 1):
            # 项目标题：设为二级标题，方便折叠展开
            p = doc.add_heading(level=2)
            title_run = p.add_run(f"{i}. {item.get('项目名称') or item.get('标题', '未命名项目')}")
            set_run_font(title_run, font_name='微软雅黑', font_size=12, bold=True)

            fields = [
                ("项目编号", "项目编号"),
                ("采购人", "采购人名称"),
                ("预算金额", "预算金额"),
                ("投标截止", "提交投标文件截止时间"),
                ("地区", "地区"),
                ("项目类型", "项目类型"),
                ("预算规模", "预算规模"),
                ("核心产品", "核心产品"),
                ("相关度评分", "相关度评分"),
                ("潜在机会", "潜在机会"),
                ("跟进建议", "跟进建议"),
                ("判断理由", "判断理由"),
                ("原文链接", "详情页链接"),
            ]

            for label, key in fields:
                val = str(item.get(key, "")).strip()
                if val:
                    info = doc.add_paragraph()
                    info.paragraph_format.left_indent = Inches(0.3)
                    info.paragraph_format.space_after = Pt(2)
                    run = info.add_run(f"· {label}：{val}")
                    set_run_font(run, font_name='微软雅黑', font_size=10.5)

            doc.add_paragraph()  # 条目间距

    # 添加各分组（仅保留高、中相关项目）
    add_section(f"高相关项目（评分 ≥ 8）— 共 {len(high)} 条", high, (192, 0, 0))
    add_section(f"中相关项目（评分 6-7）— 共 {len(medium)} 条", medium, (0, 112, 192))

    # 页脚统计
    doc.add_paragraph()
    summary = doc.add_paragraph()
    summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    summary_text = (
        f"本次共分析 {len(rows)} 条招标信息 | "
        f"高相关 {len(high)} 条 | 中相关 {len(medium)} 条"
    )
    summary_run = summary.add_run(summary_text)
    set_run_font(summary_run, font_name='微软雅黑', font_size=9)
    summary_run.font.color.rgb = RGBColor(128, 128, 128)

    doc.save(OUTPUT_DOCX)
    print(f"[√] Word 简报已生成: {OUTPUT_DOCX}")


def main():
    print("=" * 50)
    print("招标信息智能分析 Agent 启动")
    print("=" * 50)

    # 1. 分析 CSV
    analyzed_rows = analyze_csv()

    # 2. 生成 Word
    if analyzed_rows:
        generate_word_report(analyzed_rows)
    else:
        print("[!] 没有分析结果，跳过 Word 生成")

    print("\n[Done] 全部完成！")


if __name__ == "__main__":
    main()
